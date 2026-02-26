# condition_profiles/respiratory_failure.py
# Acute Respiratory Failure Condition Profile for the DRG Appeal Engine
#
# Contains all ARF-specific configuration, prompts, conditional rebuttals,
# and programmatic P/F ratio calculator.
#
# Conditional rebuttals are from Dr. Gharfeh and Dr. Bourland.
# The base engine imports this module via:
#   profile = importlib.import_module(f"condition_profiles.{CONDITION_PROFILE}")

import json
from datetime import datetime

# =============================================================================
# Profile Validation
# =============================================================================
REQUIRED_ATTRIBUTES = [
    "CONDITION_NAME", "CONDITION_DISPLAY_NAME",
    "GOLD_LETTERS_PATH", "PROPEL_DATA_PATH", "DEFAULT_TEMPLATE_PATH",
    "CLINICAL_SCORES_TABLE_NAME",
    "DENIAL_CONDITION_QUESTION", "DENIAL_CONDITION_FIELD",
    "NOTE_EXTRACTION_TARGETS", "STRUCTURED_DATA_CONTEXT", "CONFLICT_EXAMPLES",
    "WRITER_SCORING_INSTRUCTIONS",
    "ASSESSMENT_CONDITION_LABEL", "ASSESSMENT_CRITERIA_LABEL",
]


def validate_profile(profile):
    """Check that a condition profile module has all required attributes."""
    missing = [attr for attr in REQUIRED_ATTRIBUTES if not hasattr(profile, attr)]
    if missing:
        name = getattr(profile, '__name__', str(profile))
        raise AttributeError(
            f"Condition profile '{name}' is missing required attributes:\n"
            + "\n".join(f"  - {attr}" for attr in missing)
            + "\n\nSee an existing profile (e.g., sepsis.py) for the full interface."
        )

# =============================================================================
# Identity & Paths
# =============================================================================
CONDITION_NAME = "respiratory_failure"
CONDITION_DISPLAY_NAME = "Acute Respiratory Failure"
DRG_CODES = ["189", "190", "191", "207", "208"]

GOLD_LETTERS_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/gold_standard_appeals/gold_standard_appeals_arf_only"
PROPEL_DATA_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/propel_data"
DEFAULT_TEMPLATE_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/gold_standard_appeals/gold_standard_appeals_arf_only/default_respiratory_failure_template.docx"

# P/F ratio calculator — deterministic PaO2/FiO2 scoring from ABG lab data
CLINICAL_SCORES_TABLE_NAME = "fudgesicle_case_pf_scores"

# =============================================================================
# Denial Parser
# =============================================================================
DENIAL_CONDITION_QUESTION = "5. IS ACUTE RESPIRATORY FAILURE RELATED - does this denial involve acute respiratory failure, hypoxic respiratory failure, or hypercapnic respiratory failure?"
DENIAL_CONDITION_FIELD = "IS_ARF"

# =============================================================================
# Note Extraction Targets (fallback when Propel definition_summary unavailable)
# =============================================================================
NOTE_EXTRACTION_TARGETS = """## Oxygenation Status (PRIORITY - extract ALL available with timestamps)
- SpO2 readings (especially <91% on room air)
- PaO2 values (especially <60 mmHg)
- FiO2 levels and changes
- P/F ratio (PaO2/FiO2)
- Oxygen delivery method and flow rates (nasal cannula, high-flow, BiPAP, CPAP, ventilator)
- Escalation/de-escalation of respiratory support

## Ventilation Parameters
- PaCO2 values (especially >50 mmHg)
- pH values (especially <7.35)
- Respiratory rate
- Tidal volume, minute ventilation
- Ventilator settings (mode, PEEP, FiO2)

## Clinical Presentation
- Dyspnea severity and onset
- Use of accessory muscles
- Work of breathing assessments
- Mental status changes related to respiratory failure
- Cyanosis

## Treatment Response
- Response to oxygen therapy
- Response to non-invasive ventilation (BiPAP/CPAP)
- Need for intubation/mechanical ventilation
- ABG trends over time

## Underlying Etiology
- Pneumonia, COPD exacerbation, pulmonary embolism, ARDS, CHF
- Chest X-ray / CT findings
- Culture results
- Antibiotic administration"""

# =============================================================================
# Structured Data Context (for structured data extraction prompt)
# =============================================================================
DIAGNOSIS_EXAMPLES = """For example:
- "Acute respiratory failure with hypoxia"
- "Acute hypoxic respiratory failure"
- "Acute on chronic respiratory failure with hypoxia"
- "Acute hypercapnic respiratory failure"
"""

STRUCTURED_DATA_CONTEXT = """Extract a focused summary of acute respiratory failure-relevant clinical data from this timeline. Prioritize:

1. **Oxygenation Data** (with timestamps and trends):
   - SpO2 readings (all values, especially <91% on room air)
   - PaO2 values from ABGs
   - FiO2 levels
   - P/F ratio calculations
   - Oxygen delivery method changes

2. **Ventilation Data**:
   - PaCO2 values from ABGs
   - pH values
   - Respiratory rate trends
   - Ventilator settings and mode changes

3. **Respiratory Support Escalation**:
   - Timeline of oxygen delivery: room air -> nasal cannula -> high-flow -> BiPAP/CPAP -> intubation
   - Flow rates and FiO2 at each step
   - Duration on each level of support

4. **Relevant Diagnoses** (with dates - note which are pre-existing vs new):
   - Acute respiratory failure type (hypoxic, hypercapnic, mixed)
   - Underlying cause (pneumonia, COPD, PE, ARDS, CHF, etc.)"""

STRUCTURED_DATA_SYSTEM_MESSAGE = "You are a clinical data analyst specializing in respiratory failure cases."

# =============================================================================
# Conflict Detection Examples
# =============================================================================
CONFLICT_EXAMPLES = """- Note says "SpO2 maintained >92% on room air" but vitals show SpO2 <91% on room air
- Note says "patient on room air" but respiratory flowsheet shows supplemental O2
- Note says "no respiratory distress" but respiratory rate >24 documented
- Note says "ABG normal" but labs show PaO2 <60 or PaCO2 >50
- Note says "weaned to room air" but vitals show continued supplemental oxygen
- Note says "FiO2 40%" but respiratory flowsheet shows FiO2 60%"""

# =============================================================================
# Numeric Cross-Check Parameter Mapping
# =============================================================================
PARAM_TO_CATEGORY = {
    "spo2": "spo2", "sp02": "spo2", "oxygen saturation": "spo2", "o2 sat": "spo2",
    "pao2": "pao2", "po2": "pao2", "partial pressure of oxygen": "pao2",
    "paco2": "paco2", "pco2": "paco2", "partial pressure of co2": "paco2",
    "fio2": "fio2", "fi02": "fio2", "fraction of inspired oxygen": "fio2",
    "ph": "ph", "arterial ph": "ph",
    "respiratory rate": "respiratory_rate", "rr": "respiratory_rate", "resp rate": "respiratory_rate",
    "peep": "peep",
    "tidal volume": "tidal_volume",
    "lactate": "lactate", "lactic acid": "lactate",
    "creatinine": "creatinine", "cr": "creatinine",
    "gcs": "gcs", "glasgow coma scale": "gcs",
}

LAB_VITAL_MATCHERS = {
    "spo2": {
        "keywords": ["spo2", "sp02", "pulse ox", "oxygen sat"],
        "exclude": [],
        "type": "vital",
    },
    "pao2": {
        "keywords": ["po2", "pao2", "p02"],
        "exclude": ["pco2", "spco2", "venous"],
        "type": "lab",
    },
    "paco2": {
        "keywords": ["pco2", "paco2"],
        "exclude": ["spco2"],
        "type": "lab",
    },
    "fio2": {
        "keywords": ["fio2", "fi02", "fraction of inspired"],
        "exclude": [],
        "type": "lab",
    },
    "ph": {
        "keywords": ["ph"],
        "exclude": ["phosph", "pharyn", "pheno", "photo", "urine ph"],
        "type": "lab",
    },
    "respiratory_rate": {
        "keywords": ["resp", "respiratory rate", "rr"],
        "exclude": ["response"],
        "type": "vital",
    },
    "lactate": {
        "keywords": ["lactate", "lactic acid"],
        "exclude": ["dehydrogenase", "ldh"],
        "type": "lab",
    },
    "creatinine": {
        "keywords": ["creatinine"],
        "exclude": ["clearance", "urine", "ratio", "kinase", "random"],
        "type": "lab",
    },
    "gcs": {
        "keywords": ["gcs", "glasgow"],
        "exclude": ["component", "eye", "motor", "verbal"],
        "type": "vital",
    },
    "peep": {
        "keywords": ["peep"],
        "exclude": [],
        "type": "vital",
    },
    "tidal_volume": {
        "keywords": ["tidal vol"],
        "exclude": [],
        "type": "vital",
    },
}

# =============================================================================
# P/F Ratio Thresholds and Pairing Windows
# =============================================================================

# Berlin ARDS classification thresholds (P/F ratio)
# Formal Berlin definition requires PEEP >= 5 cmH2O; shown for reference regardless
PF_RENDER_THRESHOLD = 300          # Show table if any P/F < 300 (ARF-qualifying)
PF_PAIR_MAX_SECONDS = 1800         # 30 min: max gap between PaO2 and FiO2 timestamps
PEEP_PAIR_MAX_SECONDS = 7200       # 2 hours: max gap between ABG draw and PEEP reading

# =============================================================================
# Writer Prompt -- Scoring Instructions
# =============================================================================
WRITER_SCORING_INSTRUCTIONS = """5. QUANTIFY ACUTE RESPIRATORY FAILURE using diagnostic criteria from the Propel guidelines when available:
   - Reference specific values: PaO2, SpO2, PaCO2, pH, FiO2, P/F ratio
   - Hypoxic criteria: PaO2 < 60 mmHg, SpO2 < 91% on room air, P/F ratio < 300
   - Hypercapnic criteria: PaCO2 > 50 mmHg with pH < 7.35
   - Acute-on-chronic indicators: >= 10 mmHg change from baseline PaO2 or PaCO2
   - Document oxygen delivery method and escalation of respiratory support

   SPO2 THRESHOLD (CRITICAL): The hypoxic SpO2 criterion is STRICTLY LESS THAN 91%. An SpO2 of 91% is low but does NOT meet the hypoxic threshold. Do NOT cite SpO2 values of 91% or higher as evidence of hypoxemia. Only SpO2 values of 90% or below qualify.

   EVIDENCE TIMING (CRITICAL): Use clinical values from the presentation/admission timeframe ONLY. Values from 2 or more days after admission do not support an acute diagnosis at presentation. When the condition is POA, focus on the first 24 hours of admission.

   DOCUMENTATION CONSISTENCY: Do NOT use a clinical value as evidence of acute respiratory failure if the same documentation describes the patient as "stable" at that time. A value documented alongside "stable condition" cannot simultaneously support an acute diagnosis.

   P/F RATIO: If P/F ratio measurements are provided in the Computed Clinical Scores section above (from programmatic calculation), cite those deterministic values directly. Do NOT recalculate P/F ratio from raw values — use the provided measurements. P/F ratio < 300 supports ARF. P/F ratio < 200 indicates moderate-to-severe ARDS. Do NOT apply the P/F ratio criterion to acute-on-chronic respiratory failure patients. Do NOT include a P/F ratio table in the letter text — the table is rendered separately in the document.

   RESPIRATORY RATE THRESHOLD: The tachypnea threshold for ARF is respiratory rate > 20 breaths per minute (per Propel/ACDIS criteria).

   OXYGEN DELIVERY CONTEXT: When citing SpO2 readings, verify the oxygen delivery method. If the patient is on high-flow nasal cannula but FiO2 is set at 21% or less, this is essentially room air (~20% FiO2). Clarify whether the SpO2 was on room air or supplemental oxygen, as this affects the hypoxic criterion.

   SIGNS AND SYMPTOMS: Include ALL documented signs and symptoms of respiratory distress in the clinical presentation: dyspnea, difficulty breathing, increased work of breathing, use of accessory muscles, tachypnea, cyanosis, altered mental status related to respiratory compromise."""

# =============================================================================
# Assessment Labels
# =============================================================================
ASSESSMENT_CONDITION_LABEL = "acute respiratory failure DRG appeal letter"
ASSESSMENT_CRITERIA_LABEL = "PROPEL RESPIRATORY FAILURE CRITERIA (AUTHORITATIVE SOURCE - USE THIS ONLY)"

# =============================================================================
# Conditional Rebuttals (from Dr. Gharfeh and Dr. Bourland)
# =============================================================================

CONDITIONAL_REBUTTALS = [
    {
        "name": "Consecutive/Sequential SpO2 Readings",
        "trigger": "Apply ONLY if the denial requires 'consecutive,' 'sequential,' or 'back-to-back' SpO2 readings.",
        "text": """Apply ONLY if the denial requires "consecutive," "sequential," or "back-to-back" SpO2 readings.

Rebuttal points:
1. No CMS, ICD-10-CM, or Coding Clinic requirement mandates multiple sequential SpO2 readings for acute hypoxemic respiratory failure.
2. Consecutive reading requirements are proprietary payor criteria -- not nationally recognized standards.
3. A documented SpO2 below 91% on room air with clinical context and treatment response supports the diagnosis.

Cite ALL relevant timestamped low SpO2 readings below 91%, even if not consecutive.""",
    },
    {
        "name": "Persistent/Continuous Symptoms Requirement",
        "trigger": "Apply ONLY if the denial argues symptoms were not 'persistent,' 'continuous,' or 'sustained.'",
        "text": """Apply ONLY if the denial argues symptoms were not "persistent," "continuous," or "sustained."

Rebuttal points:
1. Acute respiratory failure is defined by onset and severity at presentation -- not by unchanged symptoms throughout the stay.
2. Improvement with treatment confirms appropriate intervention, not absence of the condition.
3. Per CMS/Coding Clinic, a diagnosis is reportable when documented by the provider and supported by clinical indicators and treatment at presentation.

Cite documentation of respiratory distress at presentation and interventions required.""",
    },
    # NOTE: "Proprietary Clinical Criteria (General)" rebuttal was removed per SME feedback (2026-02-23).
    # SME noted: "unfortunately they can have whatever criteria they want" — payors are allowed
    # to use their own clinical criteria, so arguing against proprietary criteria is not effective.
]

# =============================================================================
# Utility Functions (used by featurization_inference.py numeric cross-check)
# =============================================================================

def match_name(name, matcher):
    """Check if a lab/vital name matches a matcher pattern."""
    name_lower = name.lower()
    if not any(kw in name_lower for kw in matcher["keywords"]):
        return False
    if any(ex in name_lower for ex in matcher["exclude"]):
        return False
    return True


def safe_float(value):
    """Parse a numeric value, returning None if not parseable."""
    if value is None:
        return None
    try:
        cleaned = str(value).strip().replace(',', '').replace('>', '').replace('<', '')
        return float(cleaned)
    except (ValueError, TypeError):
        return None


# =============================================================================
# Clinical Scorer — P/F Ratio (Deterministic, Zero LLM Calls)
# =============================================================================
#
# PaO2 and FiO2 both come from the ABG lab panel (res_components/clarity_component).
# At Mercy, the RT records FiO2 as part of the ABG order, not in the ventilator
# flowsheet. Since both values share the same blood gas order, their timestamps
# are nearly identical — the 30-minute pairing window is a safety check.
#
# PEEP comes from the ventilator flowsheet (FLO_MEAS_IDs 1050046701, 1050056801).
# It uses a wider 2-hour pairing window since flowsheet entries have independent timing.
# =============================================================================


def _classify_berlin(ratio):
    """Return Berlin ARDS classification label for a given P/F ratio."""
    if ratio <= 100:
        return "Severe ARDS"
    elif ratio <= 200:
        return "Moderate ARDS"
    elif ratio <= 300:
        return "Mild ARDS"
    else:
        return "Normal"


def _empty_pf_result():
    """Return a zero-result dict when no P/F pairs can be calculated."""
    return {
        "pf_measurements": [],
        "worst_ratio": None,
        "worst_classification": None,
        "measurements_calculated": 0,
        # Aliases for inference.py compatibility (reads total_score, organs_scored, etc.)
        "total_score": 0,
        "organs_scored": 0,
        "organ_scores": {},
        "vasopressor_detail": [],
        "window_start": None,
        "window_end": None,
        "window_mode": "full_encounter",
    }


def calculate_clinical_scores(account_id, spark, tables,
                              admission_dt=None, poa_code=None, first_dx_timestamp=None):
    """
    Calculate P/F ratios from raw structured data tables.

    Unlike SOFA (single best 24h window), this captures ALL P/F measurements
    across the full encounter. The LLM writer uses presentation/admission values;
    the table provides complete data for CDI reviewer context.

    Zero LLM calls — purely deterministic.

    Args:
        account_id: Hospital account ID
        spark: SparkSession instance
        tables: dict with keys "labs", "vitals", "meds" -> full table names
        admission_dt: Admission datetime (stored for reference, not used for windowing)
        poa_code: POA indicator (passed in but not used — P/F captures full encounter)
        first_dx_timestamp: First ARF diagnosis timestamp (passed in but not used)

    Returns:
        dict compatible with inference.py's score loader (total_score, organs_scored,
        organ_scores, vasopressor_detail, window_start, window_end, window_mode)
        plus ARF-specific fields (pf_measurements, worst_ratio, worst_classification).
    """
    from datetime import timedelta

    print("  Calculating P/F ratios (full encounter)...")

    labs_table = tables["labs"]
    vitals_table = tables["vitals"]

    # --- Read labs (PaO2 + FiO2) ---
    try:
        labs_rows = spark.sql(f"""
            SELECT LAB_NAME, lab_value, EVENT_TIMESTAMP
            FROM {labs_table}
            WHERE lab_value IS NOT NULL AND EVENT_TIMESTAMP IS NOT NULL
            ORDER BY EVENT_TIMESTAMP ASC
        """).collect()
    except Exception as e:
        print(f"    Warning: Could not read labs: {e}")
        labs_rows = []

    # --- Read vitals (PEEP) ---
    try:
        vitals_rows = spark.sql(f"""
            SELECT VITAL_NAME, vital_value, EVENT_TIMESTAMP
            FROM {vitals_table}
            WHERE vital_value IS NOT NULL AND EVENT_TIMESTAMP IS NOT NULL
            ORDER BY EVENT_TIMESTAMP ASC
        """).collect()
    except Exception as e:
        print(f"    Warning: Could not read vitals: {e}")
        vitals_rows = []

    # --- Collect PaO2, FiO2, PEEP values ---
    pao2_values = []    # list of (value_float, timestamp)
    fio2_values = []    # list of (value_float, timestamp)
    peep_values = []    # list of (value_float, timestamp)

    for row in labs_rows:
        name = row["LAB_NAME"]
        val = safe_float(row["lab_value"])
        ts = row["EVENT_TIMESTAMP"]
        if val is None or ts is None:
            continue
        if match_name(name, LAB_VITAL_MATCHERS["pao2"]):
            pao2_values.append((val, ts))
        elif match_name(name, LAB_VITAL_MATCHERS["fio2"]):
            fio2_values.append((val, ts))

    for row in vitals_rows:
        name = row["VITAL_NAME"]
        val = safe_float(row["vital_value"])
        ts = row["EVENT_TIMESTAMP"]
        if val is None or ts is None:
            continue
        if match_name(name, LAB_VITAL_MATCHERS["peep"]):
            peep_values.append((val, ts))

    print(f"    Found: {len(pao2_values)} PaO2, {len(fio2_values)} FiO2, {len(peep_values)} PEEP values")

    if not pao2_values:
        print("    No PaO2 values found — returning empty P/F results.")
        return _empty_pf_result()

    if not fio2_values:
        print("    No FiO2 values found — returning empty P/F results.")
        return _empty_pf_result()

    # --- Pair each PaO2 with closest-in-time FiO2 (within 30-min window) ---
    pf_measurements = []

    for pao2_val, pao2_ts in pao2_values:
        # Find closest FiO2 within time window
        closest_fio2_val = None
        closest_fio2_diff = None
        for fio2_val, fio2_ts in fio2_values:
            if fio2_val <= 0:
                continue
            diff = abs((pao2_ts - fio2_ts).total_seconds())
            if diff > PF_PAIR_MAX_SECONDS:
                continue
            if closest_fio2_diff is None or diff < closest_fio2_diff:
                closest_fio2_diff = diff
                closest_fio2_val = fio2_val

        if closest_fio2_val is None:
            print(f"    Skipping PaO2={pao2_val} at {str(pao2_ts)[:19]} — no FiO2 within {PF_PAIR_MAX_SECONDS // 60} min")
            continue

        # Normalize FiO2: Epic stores as percentage (e.g., 40); ratio uses fraction (0.40)
        actual_fio2 = closest_fio2_val if closest_fio2_val <= 1.0 else closest_fio2_val / 100.0
        if actual_fio2 <= 0:
            continue

        ratio = round(pao2_val / actual_fio2, 1)
        classification = _classify_berlin(ratio)

        # Find closest PEEP reading within ±2 hours of ABG draw
        closest_peep_val = None
        closest_peep_diff = None
        for peep_val, peep_ts in peep_values:
            diff = abs((pao2_ts - peep_ts).total_seconds())
            if diff > PEEP_PAIR_MAX_SECONDS:
                continue
            if closest_peep_diff is None or diff < closest_peep_diff:
                closest_peep_diff = diff
                closest_peep_val = peep_val

        measurement = {
            "timestamp": str(pao2_ts)[:19],
            "pao2": pao2_val,
            "fio2_raw": closest_fio2_val,
            "fio2_normalized": round(actual_fio2, 4),
            "peep": closest_peep_val,
            "ratio": ratio,
            "classification": classification,
        }
        pf_measurements.append(measurement)

    if not pf_measurements:
        print("    No valid PaO2/FiO2 pairs found within time window — returning empty results.")
        return _empty_pf_result()

    # Sort by timestamp ascending
    pf_measurements.sort(key=lambda m: m["timestamp"])

    # Identify worst (lowest) ratio
    worst = min(pf_measurements, key=lambda m: m["ratio"])
    worst_ratio = worst["ratio"]
    worst_classification = worst["classification"]
    n = len(pf_measurements)

    # Build organ_scores dict indexed by string keys for JSON serialization.
    # inference.py does json.loads(row["organ_scores"]) and passes to format_scores_for_prompt.
    organ_scores_dict = {str(i): m for i, m in enumerate(pf_measurements)}

    timestamps = [m["timestamp"] for m in pf_measurements]
    window_start = timestamps[0] if timestamps else None
    window_end = timestamps[-1] if timestamps else None

    print(f"    P/F ratios calculated: {n} measurements")
    print(f"    Worst P/F ratio: {worst_ratio} ({worst_classification}) at {worst['timestamp']}")
    for m in pf_measurements:
        peep_str = f", PEEP={m['peep']}" if m["peep"] is not None else ""
        print(f"      {m['timestamp']}: PaO2={m['pao2']}, FiO2={m['fio2_raw']}{peep_str} -> P/F={m['ratio']} ({m['classification']})")

    return {
        "pf_measurements": pf_measurements,
        "worst_ratio": worst_ratio,
        "worst_classification": worst_classification,
        "measurements_calculated": n,
        # Aliases for inference.py compatibility
        "total_score": n,
        "organs_scored": n,
        "organ_scores": organ_scores_dict,
        "vasopressor_detail": [],
        "window_start": window_start,
        "window_end": window_end,
        "window_mode": "full_encounter",
    }


def write_clinical_scores_table(account_id, scores_result, spark, table_name):
    """Write P/F ratio scores to Delta case table."""
    from pyspark.sql.types import (
        StructType, StructField, StringType, IntegerType, FloatType, TimestampType
    )

    spark.sql(f"""
    CREATE TABLE IF NOT EXISTS {table_name} (
        account_id              STRING,
        total_score             INT,
        organs_scored           INT,
        organ_scores            STRING,
        vasopressor_detail      STRING,
        window_start            STRING,
        window_end              STRING,
        window_mode             STRING,
        worst_ratio             FLOAT,
        worst_classification    STRING,
        measurements_calculated INT,
        pf_measurements         STRING,
        created_at              TIMESTAMP
    ) USING DELTA
    """)

    record = [{
        "account_id":               account_id,
        "total_score":              scores_result["total_score"],
        "organs_scored":            scores_result["organs_scored"],
        "organ_scores":             json.dumps(scores_result["organ_scores"]),
        "vasopressor_detail":       json.dumps(scores_result.get("vasopressor_detail", [])),
        "window_start":             scores_result.get("window_start"),
        "window_end":               scores_result.get("window_end"),
        "window_mode":              scores_result.get("window_mode", "full_encounter"),
        "worst_ratio":              scores_result.get("worst_ratio"),
        "worst_classification":     scores_result.get("worst_classification"),
        "measurements_calculated":  scores_result.get("measurements_calculated", 0),
        "pf_measurements":          json.dumps(scores_result.get("pf_measurements", [])),
        "created_at":               datetime.now(),
    }]

    schema = StructType([
        StructField("account_id",               StringType(),    False),
        StructField("total_score",              IntegerType(),   True),
        StructField("organs_scored",            IntegerType(),   True),
        StructField("organ_scores",             StringType(),    True),
        StructField("vasopressor_detail",       StringType(),    True),
        StructField("window_start",             StringType(),    True),
        StructField("window_end",               StringType(),    True),
        StructField("window_mode",              StringType(),    True),
        StructField("worst_ratio",              FloatType(),     True),
        StructField("worst_classification",     StringType(),    True),
        StructField("measurements_calculated",  IntegerType(),   True),
        StructField("pf_measurements",          StringType(),    True),
        StructField("created_at",               TimestampType(), True),
    ])

    df = spark.createDataFrame(record, schema)
    df.write.format("delta").mode("overwrite").saveAsTable(table_name)
    print(f"  Written to {table_name}")


# =============================================================================
# Prompt Formatting and DOCX Rendering — P/F Ratio
# =============================================================================


def format_scores_for_prompt(scores_data):
    """
    Format P/F ratio measurements as a markdown table for prompt inclusion.

    Called by inference.py with case_data["clinical_scores"].
    scores_data["organ_scores"] contains indexed measurement dicts
    (e.g., {"0": {measurement}, "1": {measurement}, ...}).
    """
    if not scores_data:
        return "P/F ratio data not available. Do not reference programmatic P/F calculations."

    # Reconstruct ordered measurements from organ_scores dict
    organ_scores = scores_data.get("organ_scores", {})
    n = len(organ_scores)

    if n == 0:
        return "No valid PaO2/FiO2 pairs found. Do not reference programmatic P/F calculations."

    measurements = [organ_scores[str(i)] for i in range(n) if str(i) in organ_scores]

    # Threshold gate: only show detail if any ratio < 300 (supports ARF)
    qualifying = [m for m in measurements if m["ratio"] < PF_RENDER_THRESHOLD]
    if not qualifying:
        worst = min(measurements, key=lambda m: m["ratio"])
        return (
            f"P/F RATIO DATA: {n} measurement(s) calculated. "
            f"All ratios >= {PF_RENDER_THRESHOLD} (worst = {worst['ratio']}). "
            "No ARF-qualifying P/F values. Do not cite P/F ratio as evidence of ARF for this case."
        )

    # Find worst
    worst = min(measurements, key=lambda m: m["ratio"])

    lines = []
    lines.append(f"P/F RATIO MEASUREMENTS (calculated from raw ABG data — {n} measurement(s), full encounter):")
    lines.append(f"Worst P/F ratio: {worst['ratio']} ({worst['classification']}) at {worst['timestamp']}")
    lines.append("")
    lines.append("| Timestamp | PaO2 (mmHg) | FiO2 | PEEP (cmH2O) | P/F Ratio | Berlin Classification |")
    lines.append("|---|---|---|---|---|---|")

    for m in measurements:
        peep_display = str(m["peep"]) if m["peep"] is not None else "N/A (no vent data)"
        if m['fio2_raw'] > 1.0:
            fio2_display = f"{m['fio2_raw']}% ({m['fio2_normalized']:.2f})"
        else:
            fio2_display = f"{m['fio2_raw']} ({m['fio2_normalized'] * 100:.0f}%)"
        lines.append(
            f"| {m['timestamp']} | {m['pao2']} | {fio2_display} | {peep_display} "
            f"| {m['ratio']} | {m['classification']} |"
        )

    lines.append("")
    lines.append(
        "INSTRUCTIONS: Reference these deterministic P/F values in the appeal letter narrative. "
        "Do NOT recalculate or modify these values. Cite the worst ratio and timestamp when "
        "arguing ARF severity. Do NOT include a P/F ratio table in the letter text."
    )

    return "\n".join(lines)


def render_scores_status_note(doc, scores_data):
    """Render P/F ratio status note in the internal review section of DOCX.

    Called when the P/F table is NOT rendered (no data or no qualifying values).
    When qualifying measurements exist, this function adds nothing — the table
    in the appendix speaks for itself.
    """
    organ_scores = scores_data.get("organ_scores", {}) if scores_data else {}
    measurements = [organ_scores[str(i)] for i in range(len(organ_scores)) if str(i) in organ_scores]
    qualifying = [m for m in measurements if m["ratio"] < PF_RENDER_THRESHOLD]

    if not scores_data or len(measurements) == 0:
        note = doc.add_paragraph()
        note.add_run("P/F Ratio Table: ").bold = True
        note.add_run(
            "Not included — no valid PaO2/FiO2 pairs found in structured data for this encounter."
        )
    elif not qualifying:
        note = doc.add_paragraph()
        note.add_run("P/F Ratio Table: ").bold = True
        worst = min(measurements, key=lambda m: m["ratio"])
        note.add_run(
            f"Not included — {len(measurements)} measurement(s) calculated but all P/F ratios "
            f">= {PF_RENDER_THRESHOLD} (worst = {worst['ratio']} at {worst['timestamp']}). "
            "No ARF-qualifying P/F values."
        )


def render_scores_in_docx(doc, scores_data):
    """Render P/F ratio table in DOCX (placed after letter body as appendix)."""
    from docx.shared import Pt

    if not scores_data:
        return

    organ_scores = scores_data.get("organ_scores", {})
    measurements = [organ_scores[str(i)] for i in range(len(organ_scores)) if str(i) in organ_scores]
    qualifying = [m for m in measurements if m["ratio"] < PF_RENDER_THRESHOLD]

    if not qualifying:
        return

    n = len(measurements)
    worst = min(measurements, key=lambda m: m["ratio"])

    doc.add_paragraph()
    header = doc.add_paragraph()
    header.add_run("Appendix: P/F Ratio Summary").bold = True
    header.paragraph_format.space_after = Pt(4)

    # Encounter window note
    window_note = doc.add_paragraph()
    window_note.add_run("Data Scope: ").bold = True
    window_note.add_run(
        f"Full encounter ({n} ABG measurement(s)). "
        f"Window: {scores_data.get('window_start', 'N/A')} to {scores_data.get('window_end', 'N/A')}."
    )
    window_note.paragraph_format.space_after = Pt(4)

    # Worst ratio callout
    worst_note = doc.add_paragraph()
    worst_note.add_run("Worst P/F Ratio: ").bold = True
    worst_note.add_run(f"{worst['ratio']} ({worst['classification']}) at {worst['timestamp']}")
    worst_note.paragraph_format.space_after = Pt(6)

    # Table: header row + one row per measurement
    table = doc.add_table(rows=1 + n, cols=6, style='Table Grid')

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate([
        "Timestamp", "PaO2 (mmHg)", "FiO2", "PEEP (cmH2O)", "P/F Ratio", "Berlin Classification"
    ]):
        hdr.cells[i].text = text
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True

    # Data rows
    for row_idx, m in enumerate(measurements, start=1):
        row = table.rows[row_idx]
        peep_display = str(m["peep"]) if m["peep"] is not None else "N/A"
        if m['fio2_raw'] > 1.0:
            fio2_display = f"{m['fio2_raw']}% ({m['fio2_normalized']:.2f})"
        else:
            fio2_display = f"{m['fio2_raw']} ({m['fio2_normalized'] * 100:.0f}%)"
        row.cells[0].text = m["timestamp"]
        row.cells[1].text = str(m["pao2"])
        row.cells[2].text = fio2_display
        row.cells[3].text = peep_display
        row.cells[4].text = str(m["ratio"])
        row.cells[5].text = m["classification"]

        # Bold the ratio cell if it qualifies for ARF
        if m["ratio"] < PF_RENDER_THRESHOLD:
            for run in row.cells[4].paragraphs[0].runs:
                run.bold = True

    # Caption
    caption = doc.add_paragraph()
    caption.add_run(
        "Note: P/F ratios calculated deterministically from raw ABG data. "
        "Berlin ARDS classification shown for reference (requires PEEP >= 5 cmH2O for formal diagnosis). "
        "Bolded P/F ratio values indicate ARF-qualifying threshold (< 300)."
    )
    caption.paragraph_format.space_after = Pt(4)
