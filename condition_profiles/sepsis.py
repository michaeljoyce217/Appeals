# condition_profiles/sepsis.py
# Sepsis Condition Profile for the DRG Appeal Engine
#
# Contains all sepsis-specific configuration, prompts, scoring logic,
# and DOCX rendering. The base engine imports this module via:
#   profile = importlib.import_module(f"condition_profiles.{CONDITION_PROFILE}")

import json
import sys
from datetime import datetime

# =============================================================================
# Profile Validation
# =============================================================================
REQUIRED_ATTRIBUTES = [
    "CONDITION_NAME", "CONDITION_DISPLAY_NAME",
    "GOLD_LETTERS_PATH", "PROPEL_DATA_PATH", "DEFAULT_TEMPLATE_PATH",
    "CLINICAL_SCORES_TABLE_NAME",
    "DENIAL_CONDITION_QUESTION", "DENIAL_CONDITION_FIELD",
    "NOTE_EXTRACTION_TARGETS", "STRUCTURED_DATA_CONTEXT", "CONFLICT_EXAMPLES",
    "WRITER_SCORING_INSTRUCTIONS",
    "ASSESSMENT_CONDITION_LABEL", "ASSESSMENT_CRITERIA_LABEL",
]


def validate_profile(profile):
    """Check that a condition profile module has all required attributes."""
    missing = [attr for attr in REQUIRED_ATTRIBUTES if not hasattr(profile, attr)]
    if missing:
        name = getattr(profile, '__name__', str(profile))
        raise AttributeError(
            f"Condition profile '{name}' is missing required attributes:\n"
            + "\n".join(f"  - {attr}" for attr in missing)
            + "\n\nSee an existing profile (e.g., sepsis.py) for the full interface."
        )

# =============================================================================
# Identity & Paths
# =============================================================================
CONDITION_NAME = "sepsis"
CONDITION_DISPLAY_NAME = "Sepsis"
DRG_CODES = ["870", "871", "872"]

GOLD_LETTERS_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/gold_standard_appeals/gold_standard_appeals_sepsis_only"
PROPEL_DATA_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/propel_data"
DEFAULT_TEMPLATE_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/gold_standard_appeals/gold_standard_appeals_sepsis_only/default_sepsis_appeal_template.docx"

# Clinical scores table suffix (base engine prepends catalog + schema)
CLINICAL_SCORES_TABLE_NAME = "fudgesicle_case_sofa_scores"

# POA diagnosis filter — used by featurization_inference.py to find POA status for this condition
POA_DIAGNOSIS_FILTER = "(ICD10_CODE LIKE 'R65.2%' OR LOWER(DX_NAME) LIKE '%sepsis%')"

# =============================================================================
# Denial Parser
# =============================================================================
DENIAL_CONDITION_QUESTION = "5. IS SEPSIS RELATED - does this denial involve sepsis, severe sepsis, or septic shock?"
DENIAL_CONDITION_FIELD = "IS_SEPSIS"

# =============================================================================
# Note Extraction Targets (fallback when Propel definition_summary unavailable)
# =============================================================================
NOTE_EXTRACTION_TARGETS = """## SOFA Score Components (PRIORITY - extract ALL available)
- Respiration: PaO2/FiO2 ratio, SpO2/FiO2, oxygen requirements, ventilator settings
- Coagulation: Platelet count
- Liver: Bilirubin (total)
- Cardiovascular: MAP, hypotension, vasopressor use (drug, dose)
- CNS: GCS (Glasgow Coma Scale), mental status changes
- Renal: Creatinine, urine output

## Other Sepsis Markers
- Lactate levels (CRITICAL - include all values with times)
- WBC count, bands
- Temperature (fever, hypothermia)
- Heart rate, respiratory rate
- Blood culture results, infection source
- Antibiotic administration times

## Clinical Events
- Fluid resuscitation (volume, timing)
- ICU admission/transfer
- Physician assessments mentioning sepsis, SIRS, infection"""

# =============================================================================
# Structured Data Context (for structured data extraction prompt)
# =============================================================================
DIAGNOSIS_EXAMPLES = """For example:
- "Severe sepsis with septic shock due to Methicillin-susceptible Staphylococcus aureus"
- "Sepsis due to Escherichia coli"
"""

STRUCTURED_DATA_CONTEXT = """Extract a focused summary of sepsis-relevant clinical data from this timeline. Prioritize:

1. **SOFA Score Components** (with timestamps and trends):
   - Respiratory: PaO2/FiO2, SpO2, oxygen requirements
   - Coagulation: Platelet count
   - Liver: Bilirubin
   - Cardiovascular: MAP, vasopressor use with doses
   - CNS: GCS
   - Renal: Creatinine, urine output

2. **Sepsis Bundle Compliance**:
   - Time of suspected infection
   - Antibiotic administration (within 3 hours?)
   - Lactate measurement and remeasurement (within 6 hours if elevated?)
   - Fluid resuscitation (30 mL/kg within 3 hours if hypotensive/lactate ≥4?)
   - Vasopressor initiation (if MAP <65 after fluids?)

3. **Clinical Trajectory**:
   - When did patient meet sepsis criteria?
   - Worst values and when they occurred
   - Evidence of organ dysfunction

4. **Relevant Diagnoses** (with dates - note which are pre-existing vs new)"""

STRUCTURED_DATA_SYSTEM_MESSAGE = "You are a clinical data analyst specializing in sepsis cases."

# =============================================================================
# Conflict Detection Examples
# =============================================================================
CONFLICT_EXAMPLES = """- Note says "MAP maintained >65" but vitals show MAP values below 65
- Note says "lactate normalized" but labs show lactate still elevated (>2.0)
- Note says "no vasopressors needed" but meds show vasopressor administration
- Note says "patient alert and oriented" but GCS recorded as <15
- Note says "afebrile" but temps >38°C documented"""

# =============================================================================
# Numeric Cross-Check Parameter Mapping
# =============================================================================
PARAM_TO_CATEGORY = {
    "lactate": "lactate", "lactic acid": "lactate",
    "creatinine": "creatinine", "cr": "creatinine",
    "bilirubin": "bilirubin", "total bilirubin": "bilirubin",
    "platelets": "platelets", "platelet count": "platelets", "plt": "platelets",
    "map": "map", "mean arterial pressure": "map",
    "gcs": "gcs", "glasgow coma scale": "gcs", "glasgow": "gcs",
    "pao2": "pao2", "po2": "pao2",
    "fio2": "fio2",
}

# =============================================================================
# Writer Prompt — Scoring Instructions
# =============================================================================
WRITER_SCORING_INSTRUCTIONS = """5. SOFA SCORING:
   - If SOFA scores are provided above (total >= 2), reference them narratively when arguing organ dysfunction
   - Cite the individual organ scores and total score as clinical evidence of organ dysfunction severity
   - Do NOT include a SOFA table in the letter text — the table is rendered separately in the document
   - If no SOFA scores are provided, do not mention SOFA scoring

   RENAL TERMINOLOGY (CRITICAL): Do NOT use the term "acute kidney injury" or "AKI" unless the creatinine meets KDIGO criteria (creatinine >= 1.5x baseline). If the creatinine ratio is below 1.5x baseline, use "impaired renal function" instead. The creatinine value can still count toward the SOFA score even if it does not meet KDIGO AKI criteria. A physician may document "AKI" but the appeal letter must use the clinically supported terminology.

   POA-BASED TIMING: When citing the most severe clinical values:
   - If the sepsis diagnosis is POA "Y" (present on admission): use the most severe values within 24 hours of admission.
   - If the sepsis diagnosis is POA "N" (not present on admission): use the most severe values within 24 hours of the first sepsis documentation."""

# =============================================================================
# Assessment Labels
# =============================================================================
ASSESSMENT_CONDITION_LABEL = "sepsis DRG appeal letter"
ASSESSMENT_CRITERIA_LABEL = "PROPEL SEPSIS CRITERIA"

# =============================================================================
# Clinical Scorer — SOFA
# =============================================================================

SOFA_THRESHOLDS = {
    "respiratory": {
        0: lambda ratio: ratio >= 400,
        1: lambda ratio: 300 <= ratio < 400,
        2: lambda ratio: 200 <= ratio < 300,
        3: lambda ratio: 100 <= ratio < 200,
        4: lambda ratio: ratio < 100,
    },
    "coagulation": {
        0: lambda v: v >= 150,
        1: lambda v: 100 <= v < 150,
        2: lambda v: 50 <= v < 100,
        3: lambda v: 20 <= v < 50,
        4: lambda v: v < 20,
    },
    "liver": {
        0: lambda v: v < 1.2,
        1: lambda v: 1.2 <= v < 2.0,
        2: lambda v: 2.0 <= v < 6.0,
        3: lambda v: 6.0 <= v < 12.0,
        4: lambda v: v >= 12.0,
    },
    "cardiovascular_map": {
        0: lambda v: v >= 70,
        1: lambda v: v < 70,
    },
    "cns": {
        0: lambda v: v == 15,
        1: lambda v: 13 <= v <= 14,
        2: lambda v: 10 <= v <= 12,
        3: lambda v: 6 <= v <= 9,
        4: lambda v: v < 6,
    },
    "renal": {
        0: lambda v: v < 1.2,
        1: lambda v: 1.2 <= v < 2.0,
        2: lambda v: 2.0 <= v < 3.5,
        3: lambda v: 3.5 <= v < 5.0,
        4: lambda v: v >= 5.0,
    },
}

LAB_VITAL_MATCHERS = {
    "platelets": {
        "keywords": ["platelet"],
        "exclude": ["platelet function", "platelet antibod", "platelet aggreg"],
        "type": "lab",
    },
    "bilirubin": {
        "keywords": ["bilirubin total", "total bilirubin", "bilirubin,total", "bilirubin, total"],
        "exclude": ["direct", "indirect", "conjugated", "neonatal", "urine"],
        "type": "lab",
    },
    "creatinine": {
        "keywords": ["creatinine"],
        "exclude": ["clearance", "urine", "ratio", "kinase", "random"],
        "type": "lab",
    },
    "pao2": {
        "keywords": ["po2", "pao2", "p02"],
        "exclude": ["pco2", "spco2", "venous"],
        "type": "lab",
    },
    "fio2": {
        "keywords": ["fio2", "fi02", "fraction of inspired"],
        "exclude": [],
        "type": "lab",
    },
    "lactate": {
        "keywords": ["lactate", "lactic acid"],
        "exclude": ["dehydrogenase", "ldh"],
        "type": "lab",
    },
    "map": {
        "keywords": ["map", "mean arterial"],
        "exclude": [],
        "type": "vital",
    },
    "gcs": {
        "keywords": ["gcs", "glasgow"],
        "exclude": ["component", "eye", "motor", "verbal"],
        "type": "vital",
    },
}

VASOPRESSOR_MATCHERS = {
    "norepinephrine": ["norepinephrine", "levophed"],
    "epinephrine": ["epinephrine", "adrenaline"],
    "dopamine": ["dopamine"],
    "dobutamine": ["dobutamine"],
    "vasopressin": ["vasopressin"],
    "phenylephrine": ["phenylephrine", "neosynephrine"],
}


def match_name(name, matcher):
    """Check if a lab/vital name matches a matcher pattern."""
    name_lower = name.lower()
    if not any(kw in name_lower for kw in matcher["keywords"]):
        return False
    if any(ex in name_lower for ex in matcher["exclude"]):
        return False
    return True


def match_vasopressor(med_name):
    """Return vasopressor category if med matches, else None."""
    med_lower = med_name.lower()
    for category, keywords in VASOPRESSOR_MATCHERS.items():
        if any(kw in med_lower for kw in keywords):
            return category
    return None


def safe_float(value):
    """Parse a numeric value, returning None if not parseable."""
    if value is None:
        return None
    try:
        cleaned = str(value).strip().replace(',', '').replace('>', '').replace('<', '')
        return float(cleaned)
    except (ValueError, TypeError):
        return None


def score_organ(organ, value):
    """Score a single organ system. Returns score (0-4) or None if thresholds don't match."""
    thresholds = SOFA_THRESHOLDS.get(organ, {})
    for score in sorted(thresholds.keys(), reverse=True):
        if thresholds[score](value):
            return score
    return 0


def calculate_clinical_scores(account_id, spark, tables,
                              admission_dt=None, poa_code=None, first_dx_timestamp=None):
    """
    Calculate SOFA scores from raw structured data tables using a 24-hour window.

    Window selection follows POA-based anchoring per SME feedback:
    - POA code 1 (present on admission): fixed 24h window from admission datetime
    - POA code 2 (not present on admission): fixed 24h window from first sepsis documentation
    - No POA data / other codes: sliding window that maximizes total SOFA score

    Zero LLM calls — purely deterministic.

    Args:
        account_id: The hospital account ID
        spark: SparkSession instance
        tables: dict with keys "labs", "vitals", "meds" -> full table names
        admission_dt: Admission datetime (for POA-anchored window)
        poa_code: POA indicator (1=Y, 2=N, None=unknown)
        first_dx_timestamp: Timestamp of first sepsis diagnosis (for POA N anchoring)

    Returns:
        dict with organ_scores, total_score, organs_scored, vasopressor_detail,
        window_start, window_end, window_mode
    """
    from datetime import timedelta

    print("  Calculating SOFA scores (24h window)...")

    labs_table = tables["labs"]
    vitals_table = tables["vitals"]
    meds_table = tables["meds"]

    # --- Gather all raw values ---
    try:
        labs_rows = spark.sql(f"""
            SELECT LAB_NAME, lab_value, EVENT_TIMESTAMP
            FROM {labs_table}
            WHERE lab_value IS NOT NULL AND EVENT_TIMESTAMP IS NOT NULL
            ORDER BY EVENT_TIMESTAMP ASC
        """).collect()
    except Exception as e:
        print(f"    Warning: Could not read labs: {e}")
        labs_rows = []

    try:
        vitals_rows = spark.sql(f"""
            SELECT VITAL_NAME, vital_value, EVENT_TIMESTAMP
            FROM {vitals_table}
            WHERE vital_value IS NOT NULL AND EVENT_TIMESTAMP IS NOT NULL
            ORDER BY EVENT_TIMESTAMP ASC
        """).collect()
    except Exception as e:
        print(f"    Warning: Could not read vitals: {e}")
        vitals_rows = []

    try:
        meds_rows = spark.sql(f"""
            SELECT MED_NAME, MED_DOSE, EVENT_TIMESTAMP
            FROM {meds_table}
            WHERE MED_NAME IS NOT NULL AND EVENT_TIMESTAMP IS NOT NULL
            ORDER BY EVENT_TIMESTAMP ASC
        """).collect()
    except Exception as e:
        print(f"    Warning: Could not read meds: {e}")
        meds_rows = []

    # ---- Collect all timestamped measurements per organ ----
    # Each entry: (value, timestamp, score, raw_display)
    organ_candidates = {
        "respiratory": [],
        "coagulation": [],
        "liver": [],
        "cardiovascular": [],
        "cns": [],
        "renal": [],
    }
    all_timestamps = []

    # PaO2 and FiO2 need pairing, so collect separately first
    pao2_values = []
    fio2_values = []

    for row in labs_rows:
        name = row["LAB_NAME"]
        val = safe_float(row["lab_value"])
        ts = row["EVENT_TIMESTAMP"]
        if val is None or ts is None:
            continue

        for category, matcher in LAB_VITAL_MATCHERS.items():
            if matcher["type"] != "lab":
                continue
            if not match_name(name, matcher):
                continue

            if category == "pao2":
                pao2_values.append((val, ts))
            elif category == "fio2":
                fio2_values.append((val, ts))
            elif category == "platelets":
                score = score_organ("coagulation", val)
                organ_candidates["coagulation"].append((val, ts, score, str(val)))
                all_timestamps.append(ts)
            elif category == "bilirubin":
                score = score_organ("liver", val)
                organ_candidates["liver"].append((val, ts, score, str(val)))
                all_timestamps.append(ts)
            elif category == "creatinine":
                score = score_organ("renal", val)
                organ_candidates["renal"].append((val, ts, score, str(val)))
                all_timestamps.append(ts)

    # Build respiratory candidates from PaO2/FiO2 pairs
    for pao2_val, pao2_ts in pao2_values:
        closest_fio2 = None
        closest_diff = None
        for fio2_val, fio2_ts in fio2_values:
            if fio2_val <= 0:
                continue
            actual_fio2 = fio2_val if fio2_val <= 1.0 else fio2_val / 100.0
            if actual_fio2 <= 0:
                continue
            diff = abs((pao2_ts - fio2_ts).total_seconds())
            if closest_diff is None or diff < closest_diff:
                closest_diff = diff
                closest_fio2 = actual_fio2
        if closest_fio2 and closest_fio2 > 0:
            ratio = pao2_val / closest_fio2
            score = score_organ("respiratory", ratio)
            organ_candidates["respiratory"].append(
                (ratio, pao2_ts, score, f"PaO2/FiO2 = {round(ratio, 1)}")
            )
            all_timestamps.append(pao2_ts)

    # Vitals: MAP and GCS
    for row in vitals_rows:
        name = row["VITAL_NAME"]
        val = safe_float(row["vital_value"])
        ts = row["EVENT_TIMESTAMP"]
        if val is None or ts is None:
            continue

        if match_name(name, LAB_VITAL_MATCHERS["map"]):
            score = score_organ("cardiovascular_map", val)
            organ_candidates["cardiovascular"].append(
                (val, ts, score, f"MAP = {val}")
            )
            all_timestamps.append(ts)

        if match_name(name, LAB_VITAL_MATCHERS["gcs"]):
            if val != val:  # NaN check
                continue
            int_val = int(round(val))
            score = score_organ("cns", int_val)
            organ_candidates["cns"].append(
                (int_val, ts, score, f"GCS = {int_val}")
            )
            all_timestamps.append(ts)

    # Meds: vasopressors
    vasopressor_entries = []
    for row in meds_rows:
        med_name = row["MED_NAME"]
        if not med_name:
            continue
        category = match_vasopressor(med_name)
        if category:
            dose = safe_float(row["MED_DOSE"])
            ts = row["EVENT_TIMESTAMP"]
            if ts is None:
                continue
            vasopressor_entries.append({
                "drug": category,
                "dose": dose,
                "timestamp": ts,
            })
            all_timestamps.append(ts)

    # ---- Edge case: no timestamps at all ----
    if not all_timestamps:
        print("    No timestamped values found — returning empty scores.")
        return {
            "organ_scores": {},
            "total_score": 0,
            "organs_scored": 0,
            "vasopressor_detail": [],
            "window_start": None,
            "window_end": None,
            "window_mode": "sliding_best_score",
        }

    # ---- Determine window mode based on POA status ----
    all_timestamps = sorted(set(all_timestamps))

    # Helper: score a single 24h window
    def _score_window(w_start, w_end):
        """Score all organs within a fixed [w_start, w_end] window. Returns (organs, total, vasopressor_detail)."""
        w_organs = {}

        for organ_key, candidates in organ_candidates.items():
            if organ_key == "cardiovascular":
                continue
            best_in_window = None
            for val, ts, score, display in candidates:
                if w_start <= ts <= w_end:
                    if best_in_window is None or score > best_in_window[2]:
                        best_in_window = (val, ts, score, display)
            if best_in_window:
                w_organs[organ_key] = {
                    "score": best_in_window[2],
                    "value": best_in_window[3],
                    "timestamp": str(best_in_window[1]),
                }

        # Cardiovascular: MAP + vasopressor override within window
        map_candidates = organ_candidates["cardiovascular"]
        best_cv_score = 0
        best_cv_value = ""
        best_cv_ts = None

        for val, ts, score, display in map_candidates:
            if w_start <= ts <= w_end:
                if score > best_cv_score:
                    best_cv_score = score
                    best_cv_value = display
                    best_cv_ts = ts

        w_vasopressor_detail = []
        w_vasopressors = {}
        for entry in vasopressor_entries:
            ts = entry["timestamp"]
            if w_start <= ts <= w_end:
                w_vasopressor_detail.append({
                    "drug": entry["drug"],
                    "dose": entry["dose"],
                    "timestamp": str(ts),
                })
                drug = entry["drug"]
                dose = entry["dose"]
                if drug not in w_vasopressors:
                    w_vasopressors[drug] = (dose, ts)
                elif dose is not None:
                    existing_dose = w_vasopressors[drug][0]
                    if existing_dose is None or dose > existing_dose:
                        w_vasopressors[drug] = (dose, ts)

        for drug, (dose, ts) in w_vasopressors.items():
            cv_score = 2
            if drug == "dopamine" and dose is not None:
                if dose > 15:
                    cv_score = 4
                elif dose > 5:
                    cv_score = 3
                else:
                    cv_score = 2
            elif drug == "dobutamine":
                cv_score = 2
            elif drug in ("epinephrine", "norepinephrine") and dose is not None:
                if dose > 0.1:
                    cv_score = 4
                else:
                    cv_score = 3
            elif drug in ("epinephrine", "norepinephrine"):
                cv_score = 3

            if cv_score > best_cv_score:
                best_cv_score = cv_score
                dose_str = f" {dose}" if dose is not None else ""
                best_cv_value = f"{drug}{dose_str}"
                best_cv_ts = ts

        if best_cv_score > 0 or best_cv_ts is not None:
            w_organs["cardiovascular"] = {
                "score": best_cv_score,
                "value": best_cv_value,
                "timestamp": str(best_cv_ts) if best_cv_ts else None,
            }

        w_total = sum(o["score"] for o in w_organs.values())
        return w_organs, w_total, w_vasopressor_detail

    # Determine anchor based on POA code
    window_mode = "sliding_best_score"
    anchor_point = None

    if poa_code is not None and poa_code not in (1, 2):
        print(f"    Warning: Unexpected POA_CODE value: {poa_code} (type: {type(poa_code).__name__}) — falling back to sliding window")

    if poa_code == 1 and admission_dt is not None:
        # POA Y: anchor to admission datetime
        window_mode = "poa_anchored_admission"
        anchor_point = admission_dt
    elif poa_code == 2 and first_dx_timestamp is not None:
        # POA N: anchor to first sepsis documentation
        window_mode = "poa_anchored_first_dx"
        anchor_point = first_dx_timestamp

    if anchor_point is not None:
        # Fixed POA-anchored window
        fixed_start = anchor_point
        fixed_end = anchor_point + timedelta(hours=24)
        best_organs, best_total, best_vasopressor_detail = _score_window(fixed_start, fixed_end)
        best_window = (fixed_start, fixed_end)
        anchor_label = "admission" if window_mode == "poa_anchored_admission" else "first sepsis dx"
        print(f"    SOFA window: POA-anchored to {anchor_label} ({str(fixed_start)[:19]} — {str(fixed_end)[:19]})")
    else:
        # Sliding window: find the 24h window with highest total SOFA
        best_window = None
        best_total = -1
        best_organs = None
        best_vasopressor_detail = None

        for anchor_ts in all_timestamps:
            w_end = anchor_ts + timedelta(hours=24)
            w_organs, w_total, w_vaso = _score_window(anchor_ts, w_end)

            if w_total > best_total:
                best_total = w_total
                best_organs = w_organs
                best_window = (anchor_ts, w_end)
                best_vasopressor_detail = w_vaso

        print(f"    SOFA window: Sliding (best 24h window maximizing total score)")

    # ---- Output ----
    organs_scored = len(best_organs) if best_organs else 0
    window_start_str = str(best_window[0])[:19] if best_window else None
    window_end_str = str(best_window[1])[:19] if best_window else None

    print(f"    SOFA Total: {best_total} ({organs_scored} organs scored)")
    if best_window:
        print(f"    24h window: {window_start_str} to {window_end_str}")
    if best_organs:
        for organ, data in best_organs.items():
            print(f"      {organ}: {data['score']} ({data['value']} at {data['timestamp']})")

    return {
        "organ_scores": best_organs or {},
        "total_score": best_total if best_total >= 0 else 0,
        "organs_scored": organs_scored,
        "vasopressor_detail": best_vasopressor_detail or [],
        "window_start": window_start_str,
        "window_end": window_end_str,
        "window_mode": window_mode,
    }


def write_clinical_scores_table(account_id, scores_result, spark, table_name):
    """Write SOFA scores to case table."""
    from pyspark.sql.types import (
        StructType, StructField, StringType, IntegerType, TimestampType
    )

    spark.sql(f"""
    CREATE TABLE IF NOT EXISTS {table_name} (
        account_id STRING,
        total_score INT,
        organs_scored INT,
        organ_scores STRING,
        vasopressor_detail STRING,
        window_start STRING,
        window_end STRING,
        window_mode STRING,
        created_at TIMESTAMP
    ) USING DELTA
    """)

    record = [{
        "account_id": account_id,
        "total_score": scores_result["total_score"],
        "organs_scored": scores_result["organs_scored"],
        "organ_scores": json.dumps(scores_result["organ_scores"]),
        "vasopressor_detail": json.dumps(scores_result.get("vasopressor_detail", [])),
        "window_start": scores_result.get("window_start"),
        "window_end": scores_result.get("window_end"),
        "window_mode": scores_result.get("window_mode", "sliding_best_score"),
        "created_at": datetime.now()
    }]

    schema = StructType([
        StructField("account_id", StringType(), False),
        StructField("total_score", IntegerType(), True),
        StructField("organs_scored", IntegerType(), True),
        StructField("organ_scores", StringType(), True),
        StructField("vasopressor_detail", StringType(), True),
        StructField("window_start", StringType(), True),
        StructField("window_end", StringType(), True),
        StructField("window_mode", StringType(), True),
        StructField("created_at", TimestampType(), True)
    ])

    df = spark.createDataFrame(record, schema)
    df.write.format("delta").mode("overwrite").saveAsTable(table_name)
    print(f"  Written to {table_name}")


# =============================================================================
# DOCX Rendering — SOFA Scores
# =============================================================================

def _window_mode_label(window_mode):
    """Return human-readable label for SOFA window mode."""
    if window_mode == "poa_anchored_admission":
        return "POA-anchored to admission"
    elif window_mode == "poa_anchored_first_dx":
        return "POA-anchored to first sepsis dx"
    return "best 24h window"


ORGAN_DISPLAY_NAMES = {
    "respiratory": "Respiratory (PaO2/FiO2)",
    "coagulation": "Coagulation (Platelets)",
    "liver": "Liver (Bilirubin)",
    "cardiovascular": "Cardiovascular",
    "cns": "CNS (GCS)",
    "renal": "Renal (Creatinine)",
}

# SOFA scoring criteria descriptions for each organ system (displayed in SOFA table)
SOFA_CRITERIA = {
    "respiratory": {
        0: "PaO2/FiO2 >= 400",
        1: "PaO2/FiO2 300-399",
        2: "PaO2/FiO2 200-299",
        3: "PaO2/FiO2 100-199",
        4: "PaO2/FiO2 < 100",
    },
    "coagulation": {
        0: "Platelets >= 150 K/uL",
        1: "Platelets 100-149 K/uL",
        2: "Platelets 50-99 K/uL",
        3: "Platelets 20-49 K/uL",
        4: "Platelets < 20 K/uL",
    },
    "liver": {
        0: "Bilirubin < 1.2 mg/dL",
        1: "Bilirubin 1.2-1.9 mg/dL",
        2: "Bilirubin 2.0-5.9 mg/dL",
        3: "Bilirubin 6.0-11.9 mg/dL",
        4: "Bilirubin >= 12.0 mg/dL",
    },
    "cardiovascular": {
        0: "MAP >= 70, no vasopressors",
        1: "MAP < 70",
        2: "Dopamine <= 5 or dobutamine",
        3: "Dopamine > 5 or epi/norepi <= 0.1",
        4: "Dopamine > 15 or epi/norepi > 0.1",
    },
    "cns": {
        0: "GCS = 15",
        1: "GCS 13-14",
        2: "GCS 10-12",
        3: "GCS 6-9",
        4: "GCS < 6",
    },
    "renal": {
        0: "Creatinine < 1.2 mg/dL",
        1: "Creatinine 1.2-1.9 mg/dL",
        2: "Creatinine 2.0-3.4 mg/dL",
        3: "Creatinine 3.5-4.9 mg/dL",
        4: "Creatinine >= 5.0 mg/dL",
    },
}


def format_scores_for_prompt(scores_data):
    """Format SOFA scores as a markdown table for prompt inclusion."""
    if not scores_data or scores_data["total_score"] < 2:
        return "SOFA score < 2 or unavailable. Do not include a SOFA table in the letter."

    lines = []
    lines.append("SOFA SCORES (calculated from raw clinical data):")
    # Include the 24-hour window with mode label
    if scores_data.get("window_start") and scores_data.get("window_end"):
        mode_label = _window_mode_label(scores_data.get("window_mode", "sliding_best_score"))
        lines.append(f"24-hour scoring window ({mode_label}): {scores_data['window_start']} to {scores_data['window_end']}")
    lines.append("")
    lines.append("| Organ System | Score | Value | Criteria |")
    lines.append("|---|---|---|---|")
    for organ_key in ["respiratory", "coagulation", "liver", "cardiovascular", "cns", "renal"]:
        if organ_key in scores_data["organ_scores"]:
            data = scores_data["organ_scores"][organ_key]
            if data["score"] == 0:
                continue  # Omit zero-score rows
            display = ORGAN_DISPLAY_NAMES.get(organ_key, organ_key)
            criteria = SOFA_CRITERIA.get(organ_key, {}).get(data["score"], "")
            lines.append(f"| {display} | {data['score']} | {data['value']} | {criteria} |")
    lines.append(f"| **TOTAL** | **{scores_data['total_score']}** | {scores_data['organs_scored']} organs scored | |")
    lines.append("")

    if scores_data.get("vasopressor_detail"):
        drugs = set(v["drug"] for v in scores_data["vasopressor_detail"])
        lines.append(f"Vasopressors administered: {', '.join(sorted(drugs))}")

    return "\n".join(lines)


def render_scores_status_note(doc, scores_data):
    """Render SOFA status note in the internal review section of DOCX."""
    from docx.shared import Pt

    if not scores_data:
        sofa_note = doc.add_paragraph()
        sofa_note.add_run("SOFA Table: ").bold = True
        sofa_note.add_run("Not included — insufficient structured data to calculate SOFA scores for this encounter.")
    elif scores_data.get("total_score", 0) < 2:
        sofa_note = doc.add_paragraph()
        sofa_note.add_run("SOFA Table: ").bold = True
        sofa_note.add_run(f"Not included — total SOFA score is {scores_data['total_score']} (below threshold of 2). "
                          f"{scores_data.get('organs_scored', 0)} organ system(s) scored.")


def render_scores_in_docx(doc, scores_data):
    """Render SOFA score table in DOCX (placed after letter body)."""
    from docx.shared import Pt

    if not scores_data or scores_data.get("total_score", 0) < 2:
        return

    doc.add_paragraph()
    sofa_header = doc.add_paragraph()
    sofa_header.add_run("Appendix: SOFA Score Summary").bold = True
    sofa_header.paragraph_format.space_after = Pt(4)

    # Show the 24-hour scoring window with mode label
    if scores_data.get("window_start") and scores_data.get("window_end"):
        mode_label = _window_mode_label(scores_data.get("window_mode", "sliding_best_score"))
        window_note = doc.add_paragraph()
        window_note.add_run(f"Scoring Window ({mode_label}): ").bold = True
        window_note.add_run(f"{scores_data['window_start']} to {scores_data['window_end']}")
        window_note.paragraph_format.space_after = Pt(4)

    organ_order = ["respiratory", "coagulation", "liver", "cardiovascular", "cns", "renal"]
    # Filter to non-zero scores only
    organs_with_data = [
        o for o in organ_order
        if o in scores_data["organ_scores"] and scores_data["organ_scores"][o]["score"] > 0
    ]

    table = doc.add_table(rows=1 + len(organs_with_data) + 1, cols=4, style='Table Grid')
    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(["Organ System", "Score", "Value", "Criteria"]):
        hdr.cells[i].text = text
        for run in hdr.cells[i].paragraphs[0].runs:
            run.bold = True

    # Data rows (zero-score rows omitted)
    row_idx = 1
    for organ_key in organs_with_data:
        data = scores_data["organ_scores"][organ_key]
        display = ORGAN_DISPLAY_NAMES.get(organ_key, organ_key)
        criteria = SOFA_CRITERIA.get(organ_key, {}).get(data["score"], "")
        row = table.rows[row_idx]
        row.cells[0].text = display
        row.cells[1].text = str(data["score"])
        row.cells[2].text = str(data.get("value", ""))
        row.cells[3].text = criteria
        row_idx += 1

    # Total row
    total_row = table.rows[row_idx]
    total_row.cells[0].text = "TOTAL"
    total_row.cells[1].text = str(scores_data["total_score"])
    total_row.cells[2].text = f"{scores_data['organs_scored']} organs scored"
    total_row.cells[3].text = ""
    for cell in total_row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
